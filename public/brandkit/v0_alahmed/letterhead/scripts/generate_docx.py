from docx import Document
from docx.shared import Mm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import letterhead

# Constants
GREEN_HEX = "2f5d3a"
GOLD_HEX = "c9a24d"
GREY_HEX = "444444"

def hex_to_rgb(hex_str):
    return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))

def set_cell_bg(cell, hex_color):
    """Sets background color of a table cell."""
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), hex_color))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_table_bottom_border(table, color_hex, size=4):
    """Sets a bottom border for the entire table (green line)."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = parse_xml(r'<w:tblPr {} />'.format(nsdecls('w')))
        tbl.append(tbl_pr)
    
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = parse_xml(r'<w:tblBorders {} />'.format(nsdecls('w')))
        tbl_pr.append(tbl_borders)
    
    # rigid bottom border
    bottom = parse_xml(r'<w:bottom {} w:val="single" w:sz="{}" w:space="0" w:color="{}"/>'.format(nsdecls('w'), size, color_hex))
    # Remove existing bottom if any
    existing_bottom = tbl_borders.find(r'{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bottom')
    if existing_bottom is not None:
        tbl_borders.remove(existing_bottom)
    tbl_borders.append(bottom)

def set_cell_top_border(cell, color_hex, size=6):
    """Sets a top border for a specific cell (gold line in footer)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = parse_xml(r'<w:tcBorders {} />'.format(nsdecls('w')))
        tc_pr.append(tc_borders)
    
    top = parse_xml(r'<w:top {} w:val="single" w:sz="{}" w:space="0" w:color="{}"/>'.format(nsdecls('w'), size, color_hex))
    tc_borders.append(top)

def create_docx():
    doc = Document()
    
    # --------------------------------------------------
    # Page Setup (Matching letterhead.py geometry)
    # --------------------------------------------------
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    
    # PDF uses margin_x = 20mm
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.top_margin = Mm(10) # Minimize top margin to push header up
    section.bottom_margin = Mm(22)
    section.header_distance = Mm(10)
    section.footer_distance = Mm(10)
    
    # --------------------------------------------------
    # Header
    # --------------------------------------------------
    header = section.header
    
    # Layout: [Logo (35mm)] [Spacer (Auto)] [Wordmark+Tagline (90mm)]
    # Total Width = 210 - 20 - 20 = 170mm
    # Logo = 35mm
    # Wordmark = 90mm
    # Spacer = 170 - 35 - 90 = 45mm
    
    htable = header.add_table(rows=1, cols=3, width=Mm(170))
    htable.autofit = False
    
    # Column 1: Logo
    htable.columns[0].width = Mm(35)
    cell_logo = htable.cell(0, 0)
    cell_logo.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM # Align bottom to match baseline of wordmark area
    p_logo = cell_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_logo = p_logo.add_run()
    run_logo.add_picture(letterhead.LOGO, width=Mm(35))
    
    # Column 2: Spacer
    htable.columns[1].width = Mm(45)
    cell_spacer = htable.cell(0, 1)
    
    # Column 3: Wordmark & Tagline
    htable.columns[2].width = Mm(90)
    cell_right = htable.cell(0, 2)
    cell_right.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    
    # --- Wordmark ---
    # Clear default
    cell_right.paragraphs[0].clear()
    p_wm = cell_right.add_paragraph()
    p_wm.alignment = WD_ALIGN_PARAGRAPH.CENTER # Centered in this 90mm cell
    run_wm = p_wm.add_run()
    run_wm.add_picture(letterhead.WORDMARK, width=Mm(90))
    
    # --- Tagline ---
    p_tag = cell_right.add_paragraph(letterhead.TAGLINE)
    p_tag.alignment = WD_ALIGN_PARAGRAPH.CENTER # Centered in this 90mm cell
    p_tag.space_before = Pt(2) # Slight gap
    p_tag.space_after = Pt(2)  # Gap before the green line
    r_tag = p_tag.runs[0]
    r_tag.font.name = 'Arial'
    r_tag.font.size = Pt(11)
    r_tag.font.color.rgb = hex_to_rgb(GOLD_HEX)
    
    # --- Separator Line (Green) ---
    # Apply to the whole table
    set_table_bottom_border(htable, GREEN_HEX, size=6) # size 6 = 0.75pt approx (1/8 units)
    
    # --------------------------------------------------
    # Footer
    # --------------------------------------------------
    footer = section.footer
    
    # Table for Green Background
    ftable = footer.add_table(rows=1, cols=1, width=Mm(170))
    ftable.autofit = False
    ftable.columns[0].width = Mm(170)
    
    fcell = ftable.cell(0, 0)
    set_cell_bg(fcell, GREEN_HEX)
    # Add Gold Top Border to this cell
    set_cell_top_border(fcell, GOLD_HEX, size=12) # Thicker line
    
    fcell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # --- Contact Line ---
    fcell.paragraphs[0].clear()
    p_contact = fcell.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.space_before = Pt(8)
    r_c = p_contact.add_run(letterhead.CONTACT_LINE)
    r_c.font.name = 'Arial'
    r_c.font.bold = True
    r_c.font.size = Pt(9)
    r_c.font.color.rgb = hex_to_rgb(GOLD_HEX)
    
    # --- Address Line ---
    # "Al Ahmad Continental" text is in PDF footer? No, just address.
    # Reconstructing address
    address_text_pre = "#579, 32nd 'D' Cross, 10th Main Road, 4th Block, Jayanagar, Bangalore - 560011, Karnataka, "
    
    p_addr = fcell.add_paragraph()
    p_addr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_addr.space_after = Pt(8)
    
    r_a1 = p_addr.add_run(address_text_pre)
    r_a1.font.name = 'Arial'
    r_a1.font.size = Pt(9)
    r_a1.font.color.rgb = hex_to_rgb("FFFFFF")
    
    r_a2 = p_addr.add_run("India")
    r_a2.font.name = 'Arial'
    r_a2.font.size = Pt(9)
    r_a2.font.bold = True
    r_a2.font.color.rgb = hex_to_rgb("FFFFFF")
    
    r_a3 = p_addr.add_run(".")
    r_a3.font.name = 'Arial'
    r_a3.font.size = Pt(9)
    r_a3.font.color.rgb = hex_to_rgb("FFFFFF")

    # --------------------------------------------------
    # Body Content (Placeholder)
    # --------------------------------------------------
    doc.add_paragraph("\n")
    p_body = doc.add_paragraph("Dear Sir/Madam,")
    p_body.paragraph_format.space_after = Pt(12)
    doc.add_paragraph("We are pleased to submit our proposal...")

    # Output
    output_filename = "Al_Ahmad_Continental_Letterhead.docx"
    doc.save(output_filename)
    print(f"✅ Generated DOCX: {output_filename}")

if __name__ == "__main__":
    create_docx()